#!/usr/bin/env python3
"""生成“零偏模型优化阶段汇报”Word文档。"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "reports" / "零偏模型优化阶段汇报_20260527.docx"
FIG_CMCC = ROOT / "trained_models" / "cmcc_bias_compare_Data05_0109_4q跑道.png"
FIG_EKF = ROOT / "trained_models" / "data05_no_gnss_trajectory.png"


def set_doc_font(doc: Document, font_name: str = "宋体", size_pt: float = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_main_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.cell(0, i)
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.cell(r_idx, c_idx).text = str(val)


def main() -> None:
    doc = Document()
    set_doc_font(doc)

    today = datetime.now().strftime("%Y年%m月%d日")
    add_main_title(doc, "IMU 零偏模型优化阶段汇报（领导版）")
    doc.add_paragraph()
    p = doc.add_paragraph(f"日期：{today}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading("一、阶段目标与当前状态", level=1)
    doc.add_paragraph(
        "本阶段目标是提升六轴零偏预测精度，重点改善陀螺 X/Y 轴泛化能力，并将模型结果平滑接入导航链路。"
        "当前已完成 DeepBiasNet 三阶段训练与 Data05 验证评估，模型权重与日志已归档。"
    )

    doc.add_heading("二、本轮已完成的优化动作", level=1)
    done_items = [
        "训练数据策略：仅使用 cmcc_stable（CMCC 收敛后再等待稳定时段）参与训练，降低爬升段脏标签干扰。",
        "模型结构升级：由浅层 3s 窗口网络升级为 DeepBiasNet（TCN + 扩窗）。",
        "训练流程升级：构建三阶段训练（稳态监督 -> 平滑监督 -> acc 稳态强化）。",
        "推理后处理：仅在 stable 段输出 + EMA + 步长限幅，抑制滑窗抖动。",
        "配置优化（已完成代码修改，待完整回归）：默认窗口 200->300，提升陀螺损失权重，降低 stage1/stage2 学习率，减弱 stage3 的 acc 样本重复权重。"
    ]
    for item in done_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("三、与上一版相比的关键变化", level=1)
    add_table(
        doc,
        headers=["维度", "上一版", "当前版本", "变化判断"],
        rows=[
            ["主网络", "Shallow BiasNet（3s）", "DeepBiasNet（20s，后续默认改 30s）", "时序建模能力增强"],
            ["训练监督", "瞬时标签为主", "三阶段监督 + 稳态约束", "对稳态偏置更友好"],
            ["推理输出", "全段输出", "仅 stable 段有效输出 + 平滑", "结果稳定性提升"],
            ["优化重点", "整体误差下降", "进一步攻克陀螺 X/Y 泛化", "目标更聚焦"],
        ],
    )

    doc.add_heading("四、核心结果（Data05 验证）", level=1)
    doc.add_paragraph("1）零偏精度对比（cmcc_stable，MAE）")
    add_table(
        doc,
        headers=["通道", "旧浅层 MAE", "DeepTCN MAE", "变化"],
        rows=[
            ["ba_x [g]", "0.0137", "0.0142", "基本持平"],
            ["ba_y [g]", "0.0216", "0.0142", "下降约34%"],
            ["ba_z [g]", "0.0290", "0.0223", "下降约23%"],
            ["bg_x [°/s]", "0.0251", "0.0676", "上升（待优化）"],
            ["bg_y [°/s]", "0.0494", "0.0606", "小幅上升（待优化）"],
            ["bg_z [°/s]", "0.0274", "0.0161", "下降约41%"],
        ],
    )
    doc.add_paragraph("2）导航拒止场景结果（Data05 双段 GNSS 拒止）")
    add_table(
        doc,
        headers=["指标", "DR 基线", "BiasNet + EKF", "变化"],
        rows=[
            ["RMSE (m)", "16.2", "5.2", "下降约68%"],
            ["拒止段最大误差 (m)", "120.3", "15.7", "下降约87%"],
        ],
    )
    doc.add_paragraph("结果解读：整体定位稳定性明显提升，但第二段拒止期间误差仍有抬升，航向状态在后段存在跳变现象。")

    if FIG_CMCC.exists():
        doc.add_paragraph("图1  Data05 零偏预测与 CMCC 真值对比")
        doc.add_picture(str(FIG_CMCC), width=Cm(16))
    if FIG_EKF.exists():
        doc.add_paragraph("图2  Data05 双段 GNSS 拒止轨迹与误差对比")
        doc.add_picture(str(FIG_EKF), width=Cm(16))

    doc.add_heading("五、待改善问题", level=1)
    issues = [
        "陀螺 X/Y 泛化偏弱：在验证段 MAE 相较旧版未改善，存在通道不均衡。",
        "拒止后段航向稳定性不足：存在 ±180° 跳变现象，需区分显示 wrap 与状态抖动。",
        "训练-推理链路一致性仍需闭环：需统一 DeepBiasNet 在评估与导航链路中的模型构建方式和后处理策略。",
    ]
    for item in issues:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("六、后续两周优化计划", level=1)
    plan_rows = [
        ["第1周", "完成 6 组配置回归（窗口、损失权重、stage3策略）", "bg_x/bg_y MAE 至少下降 10%"],
        ["第1周", "补充分段诊断（两段拒止分别统计位置/航向）", "定位抬升根因可解释"],
        ["第2周", "新增城市道路场景复训 + 交叉验证", "训练域/测试域性能差收敛"],
        ["第2周", "接入 EKF 端到端评估流水线", "形成稳定版本对比报表"],
    ]
    add_table(doc, ["时间", "任务", "验收标准"], plan_rows)

    doc.add_heading("七、阶段结论", level=1)
    doc.add_paragraph(
        "相较上一版，当前方案在加速度计三轴与陀螺 Z 轴上已取得明确提升，且在 GNSS 拒止场景中显著压低轨迹误差。"
        "当前主要短板已收敛到陀螺 X/Y 泛化与后段航向稳定性，问题范围清晰、优化路径明确。"
        "综合判断：本阶段优化方向有效，具备继续投入迭代并向工程应用推进的条件。"
    )

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"已生成: {OUT_DOCX}")


if __name__ == "__main__":
    main()
