#!/usr/bin/env python3
"""生成 CMCC 零偏模型汇报 Word 文档。"""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
EVAL_JSON = ROOT / "training_logs" / "cmcc_eval_20260526_233306.json"
INFO_JSON = ROOT / "trained_models" / "biasnet_info_cmcc.json"
OUT_DOCX = ROOT / "reports" / "CMCC零偏模型汇报_20260527.docx"
FIG_DIR = ROOT / "trained_models"

CHANNEL_CN = {
    "ba_x_g": "加速度计 X 零偏 (g)",
    "ba_y_g": "加速度计 Y 零偏 (g)",
    "ba_z_g": "加速度计 Z 零偏 (g)",
    "bg_x_degs": "陀螺仪 X 零偏 (°/s)",
    "bg_y_degs": "陀螺仪 Y 零偏 (°/s)",
    "bg_z_degs": "陀螺仪 Z 零偏 (°/s)",
}

VAL_SEGMENT_ID = "Data05"
VAL_FIG = ("Data05_0109_4圈跑道", "cmcc_bias_compare_Data05_0109_4q跑道.png")


def set_doc_font(doc: Document, font_name: str = "宋体", size_pt: float = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_title(doc: Document, text: str, level: int = 0) -> None:
    if level == 0:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(18)
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    else:
        doc.add_heading(text, level=level)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)


def main() -> None:
    eval_data = json.loads(EVAL_JSON.read_text(encoding="utf-8"))
    info = json.loads(INFO_JSON.read_text(encoding="utf-8"))

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_font(doc)

    today = datetime.now().strftime("%Y年%m月%d日")
    add_title(doc, "IMU 六轴零偏估计模型（BiasNet）\n阶段性成果汇报")
    doc.add_paragraph()
    p = doc.add_paragraph(f"汇报日期：{today}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(
        "说明：本报告基于 Data0109 实车数据，以 CMCC 组合导航输出的六轴零偏为监督真值，"
        "训练深度学习网络在线估计加速度计与陀螺仪零偏，用于后续 EKF/组合导航中的 IMU 误差补偿。"
    )

    # --- 1 目标 ---
    add_title(doc, "一、工作目标", 1)
    doc.add_paragraph(
        "在车辆动态行驶过程中，实时估计 IMU 六轴零偏（3 轴加速度计偏置 + 3 轴陀螺仪偏置），"
        "使估计值在 CMCC 标定收敛后的稳态段内接近 CMCC 真值，为无 GNSS 或弱 GNSS 场景下的惯导精度提供支撑。"
    )
    doc.add_paragraph("输出通道（6 维）：")
    for ch in info["output_channels"]:
        doc.add_paragraph(ch, style="List Bullet")

    # --- 2 数据 ---
    add_title(doc, "二、数据与划分", 1)
    doc.add_paragraph("数据来源：Data0109 实车采集（2026-01-09），采样对齐 10 Hz。")
    doc.add_paragraph(
        "训练段：Data01～Data04（四段不同工况：8 字、跑道组合）。"
    )
    doc.add_paragraph(
        f"验证段：{info['val_segment']}（全程未参与梯度更新，仅用于早停与泛化评估）。"
    )
    doc.add_paragraph(
        "有效学习区间：仅使用 CMCC 标定类型≥3 且收敛后再等待 60 s 的 cmcc_stable 稳态段；"
        "标定爬升期与起步静止段不参与训练，避免错误标签。"
    )

    # --- 3 模型 ---
    add_title(doc, "三、模型结构", 1)
    arch = "DeepBiasNet（深度时序卷积）" if info.get("arch") == "deep" else "BiasNet（浅层）"
    doc.add_paragraph(f"架构：{arch}")
    doc.add_paragraph(f"输入滑窗：{info['window_size']} 帧（{info['window_size'] * info['target_dt']:.0f} s @ {info['target_dt']} s/帧）")
    doc.add_paragraph(
        "网络要点：因果 TCN 残差结构，膨胀率 1/2/4/8/16/32，全局池化后输出 6 维零偏；"
        "参数量约 9.1 万。推理阶段对加速度计输出施加 ±0.25 g 物理限幅，并配合 EMA 平滑与步长限幅后处理。"
    )
    doc.add_paragraph("模型文件（工程目录 trained_models/）：")
    files = [
        ("最终权重", "biasnet_weights_cmcc.weights.h5"),
        ("阶段一权重", "biasnet_weights_cmcc_s1.weights.h5"),
        ("训练元数据", "biasnet_info_cmcc.json"),
        ("加速度计推理校准", "biasnet_cmcc_acc_calib.json"),
    ]
    for name, fn in files:
        fp = ROOT / "trained_models" / fn
        size_kb = fp.stat().st_size / 1024 if fp.exists() else 0
        doc.add_paragraph(f"{name}：{fn}（{size_kb:.0f} KB）", style="List Bullet")

    # --- 4 训练 ---
    add_title(doc, "四、训练策略", 1)
    doc.add_paragraph("三阶段微调（监督信号均为 CMCC 稳态段）：")
    stages = [
        ("阶段一", info["stage1"], "cmcc_stable 全通道 Huber 损失，建立粗对齐"),
        ("阶段二", info["stage2"], "标签平滑 + 全窗 stable 约束，细化六通道"),
        ("阶段三", info["stage3"], "加速度计段末 60 s 均值强化 + 陀螺 Huber，突出稳态尾段"),
    ]
    stage_rows = []
    for label, st, desc in stages:
        mae = st.get("best_val_mae")
        mae_s = f"{mae:.4f}" if mae is not None else "—"
        stage_rows.append([label, desc, f"验证集最佳 MAE：{mae_s}"])
    add_table(doc, ["阶段", "策略说明", "验证指标"], stage_rows)

    # --- 5 结果（仅验证集）---
    add_title(doc, "五、验证集评估结果", 1)
    doc.add_paragraph(
        "汇报指标仅给出验证集 Data05（未参与训练）。Data01～Data04 为训练集，"
        "其拟合误差通常低于真实泛化水平，不作为对外汇报依据。"
    )
    doc.add_paragraph(
        f"评估时间戳：{eval_data['timestamp']}；评估掩码：{eval_data['report_mask']}；"
        f"指标为预测值与 CMCC 真值的平均绝对误差 MAE。"
    )

    val = next(r for r in eval_data["results"] if r.get("id") == VAL_SEGMENT_ID)
    pc = val["cmcc_stable"]["per_channel"]
    headers = ["通道", "MAE", "RMSE", "单位"]
    result_rows = []
    for key in ["ba_x_g", "ba_y_g", "ba_z_g", "bg_x_degs", "bg_y_degs", "bg_z_degs"]:
        unit = "g" if key.startswith("ba_") else "°/s"
        result_rows.append([
            CHANNEL_CN[key],
            f"{pc[key]['mae']:.4f}",
            f"{pc[key]['rmse']:.4f}",
            unit,
        ])
    add_table(doc, headers, result_rows)

    doc.add_paragraph()
    doc.add_paragraph(
        "说明：加速度计三轴 MAE 约 0.018～0.036 g，表现较好；陀螺仪 Z 轴 MAE 约 0.019 °/s；"
        "陀螺仪 X/Y 轴 MAE 约 0.05 °/s，相对偏大，为后续优化方向。"
    )
    doc.add_paragraph(
        f"Data05 稳态样本数：{val['cmcc_stable']['per_channel']['ba_x_g']['n']} 帧；"
        f"CMCC 可用比例：{val['cmcc_ok_ratio']*100:.1f}%；"
        f"稳态可用比例：{val['cmcc_stable_ratio']*100:.1f}%。"
    )

    # --- 6 图表（仅验证集）---
    add_title(doc, "六、验证集预测与真值对比", 1)
    doc.add_paragraph(
        f"验证段：{VAL_FIG[0]}。绿色底为 cmcc_stable 稳态段；蓝线为 CMCC 真值，"
        "橙线为网络预测（经后处理）；仅稳态段输出有效预测。"
    )
    fig_path = FIG_DIR / VAL_FIG[1]
    if fig_path.exists():
        doc.add_picture(str(fig_path), width=Cm(16))

    # --- 7 结论 ---
    add_title(doc, "七、阶段结论与后续计划", 1)
    doc.add_paragraph(
        "已完成基于 CMCC 真值监督的 DeepBiasNet 三阶段训练，验证集六轴零偏估计达到厘米级加速度计偏置、"
        "百分度级陀螺仪偏置量级（详见上表）。模型权重与评估日志已归档，可接入 EKF 导航链路做端到端验证。"
    )
    doc.add_paragraph("后续建议：", style="List Bullet")
    for item in [
        "将当前 DeepBiasNet 权重接入 ekf_navigator 推理链路，对比纯 CMCC 与网络补偿后的轨迹误差；",
        "针对陀螺 X/Y 泛化偏弱，尝试加大陀螺损失权重或引入更长时序模型；",
        "在新增城市道路等场景数据上复训并交叉验证。",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    doc.save(OUT_DOCX)
    print(f"已生成: {OUT_DOCX}")


if __name__ == "__main__":
    main()
