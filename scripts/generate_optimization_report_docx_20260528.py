#!/usr/bin/env python3
"""生成 2026-05-28 零偏模型优化阶段汇报 Word 文档。"""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parents[1]
LOG_DIR = ROOT / "training_logs"
INFO_JSON = ROOT / "trained_models" / "biasnet_info_cmcc.json"
OUT_DOCX = ROOT / "reports" / "零偏模型优化阶段汇报_20260528.docx"

EVAL_BASELINE = LOG_DIR / "cmcc_eval_20260527_193507.json"
EVAL_TC_FAIL = LOG_DIR / "cmcc_eval_20260528_150140.json"
EVAL_TC_OK = LOG_DIR / "cmcc_eval_20260528_153136.json"
EVAL_BGY_HIGH = LOG_DIR / "cmcc_eval_20260528_163047.json"
EVAL_BEST = LOG_DIR / "cmcc_eval_20260528_165527.json"

FIG_DATA05 = ROOT / "trained_models" / (
    "cmcc_bias_compare_20260528_165527_Data05_0109_4q跑道.png"
)
FIG_EKF = ROOT / "trained_models" / "data05_no_gnss_trajectory.png"

CHANNEL_CN = {
    "ba_x_g": "加速度计 X 零偏",
    "ba_y_g": "加速度计 Y 零偏",
    "ba_z_g": "加速度计 Z 零偏",
    "bg_x_degs": "陀螺仪 X 零偏",
    "bg_y_degs": "陀螺仪 Y 零偏",
    "bg_z_degs": "陀螺仪 Z 零偏",
}
CHANNEL_UNIT = {
    "ba_x_g": "g",
    "ba_y_g": "g",
    "ba_z_g": "g",
    "bg_x_degs": "°/s",
    "bg_y_degs": "°/s",
    "bg_z_degs": "°/s",
}


def set_doc_font(doc: Document, font_name: str = "宋体", size_pt: float = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_main_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.cell(0, i)
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.cell(r_idx, c_idx).text = str(val)


def load_eval(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def data05_mae(eval_data: dict) -> dict[str, float]:
    val = next(r for r in eval_data["results"] if r["id"] == "Data05")
    pc = val["metrics"]["per_channel"]
    return {k: float(pc[k]["mae"]) for k in CHANNEL_CN}


def fmt_delta(new: float, old: float) -> str:
    d = new - old
    if abs(d) < 0.0005:
        return "基本持平"
    pct = abs(d) / max(abs(old), 1e-6) * 100
    if d < 0:
        return f"下降约{pct:.0f}%"
    return f"上升约{pct:.0f}%"


def main() -> None:
    info = json.loads(INFO_JSON.read_text(encoding="utf-8")) if INFO_JSON.exists() else {}
    m_base = data05_mae(load_eval(EVAL_BASELINE))
    m_tc_fail = data05_mae(load_eval(EVAL_TC_FAIL))
    m_tc_ok = data05_mae(load_eval(EVAL_TC_OK))
    m_bgy_high = data05_mae(load_eval(EVAL_BGY_HIGH))
    m_best = data05_mae(load_eval(EVAL_BEST))
    eval_best = load_eval(EVAL_BEST)

    doc = Document()
    set_doc_font(doc)

    report_date = "2026年05月28日"
    add_main_title(doc, "IMU 零偏模型优化阶段汇报")
    doc.add_paragraph()
    p = doc.add_paragraph(f"汇报日期：{report_date}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph(
        "说明：本报告汇总 2026-05-28 当日对 DeepBiasNet（CMCC 监督）的训练配置优化、"
        "时间一致性损失迭代、陀螺通道权重调参过程，以及当前推荐最优模型在验证集 Data05 上的评估结果。"
        "评估掩码均为 cmcc_stable（CMCC 标定收敛后再等待 60 s 的稳态段）。"
    )

    # --- 1 ---
    doc.add_heading("一、阶段目标", level=1)
    doc.add_paragraph(
        "在 Data0109 实车数据上，以 CMCC 六轴零偏为监督真值，提升 BiasNet 在线零偏估计精度，"
        "重点解决陀螺 X/Y 泛化偏弱、预测曲线抖动偏大等问题，并形成可接入 EKF 导航链路的稳定权重版本。"
    )

    # --- 2 ---
    doc.add_heading("二、当日优化动作汇总", level=1)
    actions = [
        "训练窗口：Deep 默认窗口 200 → 300 帧（30 s @ 10 Hz），增强长时序建模能力。",
        "损失加权：阶段1/2 六轴 Huber 改为 acc/gyro 分权重（提高陀螺通道梯度占比）。",
        "学习率：deep 阶段1/2 学习率下调（5e-4→3e-4，1e-4→5e-5），减轻验证集过拟合风险。",
        "阶段3：acc_sample_dup 2→1，减弱加速度计任务对六轴梯度的主导。",
        "时间一致性：阶段2 引入弱时间一致性正则（λ=0.008，仅约束 bg_x/bg_y，带标签差分门控）。",
        "陀螺权重调参：LOSS_GYRO_WEIGHTS 经历 [2.0,2.0,1.5] → [2.0,2.4,1.5] → [2.0,2.25,1.9] 三轮对比。",
    ]
    for item in actions:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("三、当前推荐训练配置（最优版）", level=1)
    add_table(
        doc,
        headers=["配置项", "取值", "说明"],
        rows=[
            ["网络", "DeepBiasNet", "TCN 残差，window=300"],
            ["训练掩码", "cmcc_stable", "cmcc_ok 后再等待 60 s"],
            ["LOSS_GYRO_WEIGHTS", "[2.0, 2.25, 1.9]", "bg_x / bg_y / bg_z 通道权重"],
            ["时间一致性", "λ=0.008，仅 bg_x/bg_y", "标签差分门控 ≤0.03 °/s"],
            ["stage1 LR (deep)", "3e-4", "—"],
            ["stage2 LR (deep)", "5e-5", "含时间一致性损失"],
            ["stage3 acc_sample_dup", "1", "acc 末 60 s 常值标签"],
            ["评估日志", EVAL_BEST.name, "2026-05-28 16:55"],
        ],
    )

    # --- 4 ---
    doc.add_heading("四、调参过程与 Data05 验证对比（MAE）", level=1)
    doc.add_paragraph(
        "下表为验证集 Data05、cmcc_stable 掩码下的六轴 MAE。"
        "目标通道 bg_y 期望 ≤0.043 °/s。"
    )

    compare_rows = []
    for key in CHANNEL_CN:
        unit = CHANNEL_UNIT[key]
        compare_rows.append([
            f"{CHANNEL_CN[key]} ({unit})",
            f"{m_base[key]:.4f}",
            f"{m_tc_ok[key]:.4f}",
            f"{m_bgy_high[key]:.4f}",
            f"{m_best[key]:.4f}",
        ])
    add_table(
        doc,
        headers=["通道", "基线(05-27)", "时间一致性修正(153136)", "bg_y=2.4(163047)", "折中权重(165527)"],
        rows=compare_rows,
    )

    doc.add_paragraph()
    doc.add_paragraph("相对上一版稳定结果（153136）的变化：")
    delta_rows = []
    for key in CHANNEL_CN:
        delta_rows.append([
            CHANNEL_CN[key],
            f"{m_tc_ok[key]:.4f} → {m_best[key]:.4f}",
            fmt_delta(m_best[key], m_tc_ok[key]),
        ])
    add_table(doc, ["通道", "153136 → 165527", "变化"], delta_rows)

    doc.add_paragraph()
    doc.add_paragraph("过程结论：")
    process_notes = [
        "首次时间一致性（150140，λ=0.05，约束三轴陀螺）导致 bg_z 严重退化（0.0377→0.0831），不可采用。",
        "二次修正（153136）显著改善 bg_z/bg_x，但 bg_y 仍偏高（0.0455）。",
        "单独提高 bg_y 权重至 2.4（163047）可使 bg_y 达标（0.0284），但 bg_z 劣化至 0.0724。",
        "折中权重 [2.0, 2.25, 1.9]（165527）同时实现 bg_y≤0.043 与 bg_z 最优（0.0242），为当前推荐版本。",
    ]
    for note in process_notes:
        doc.add_paragraph(note, style="List Bullet")

    # --- 5 ---
    doc.add_heading("五、当前最优模型验证结果（Data05）", level=1)
    val = next(r for r in eval_best["results"] if r["id"] == "Data05")
    pc = val["metrics"]["per_channel"]
    result_rows = []
    for key in CHANNEL_CN:
        result_rows.append([
            CHANNEL_CN[key],
            f"{pc[key]['mae']:.4f}",
            f"{pc[key]['rmse']:.4f}",
            CHANNEL_UNIT[key],
        ])
    add_table(doc, ["通道", "MAE", "RMSE", "单位"], result_rows)

    doc.add_paragraph()
    doc.add_paragraph(
        f"评估时间戳：{eval_best['timestamp']}；"
        f"稳态样本数：{pc['ba_x_g']['n']} 帧；"
        f"稳态可用比例：{val['cmcc_stable_ratio']*100:.1f}%。"
    )
    doc.add_paragraph(
        "末 60 s 稳态均值偏差（pred_mean - cmcc_mean）："
        f"bg_y = {val['metrics']['tail_60s']['bg_y_degs']['mean_err']:+.4f} °/s，"
        f"bg_z = {val['metrics']['tail_60s']['bg_z_degs']['mean_err']:+.4f} °/s。"
    )

    if FIG_DATA05.exists():
        doc.add_heading("六、验证集对比图（Data05）", level=1)
        doc.add_paragraph(
            "验证段：Data05_0109_4圈跑道。绿色底为 cmcc_stable；黑线为 CMCC 真值，"
            "蓝线为 BiasNet 预测（仅稳态段输出，经 EMA + 步长限幅后处理）。"
        )
        doc.add_picture(str(FIG_DATA05), width=Cm(16))

    if FIG_EKF.exists():
        doc.add_heading("七、导航拒止场景参考（Data05）", level=1)
        doc.add_paragraph(
            "双段 GNSS 拒止场景下，BiasNet + EKF 相对纯 DR 的定位误差显著降低"
            "（RMSE 约 5.2 m，拒止段最大误差约 15.7 m，参考前期验证）。"
        )
        doc.add_picture(str(FIG_EKF), width=Cm(16))

    # --- 8 ---
    doc.add_heading("八、产出文件", level=1)
    outputs = [
        ("推荐权重", ROOT / "trained_models/biasnet_weights_cmcc.weights.h5"),
        ("训练元数据", ROOT / "trained_models/biasnet_info_cmcc.json"),
        ("acc 校准", ROOT / "trained_models/biasnet_cmcc_acc_calib.json"),
        ("评估日志", EVAL_BEST),
        ("对比图", FIG_DATA05),
    ]
    for name, fp in outputs:
        size_kb = fp.stat().st_size / 1024 if fp.exists() else 0
        rel = fp.relative_to(ROOT) if fp.is_relative_to(ROOT) else fp.name
        doc.add_paragraph(f"{name}：{rel}（{size_kb:.0f} KB）", style="List Bullet")

    # --- 9 ---
    doc.add_heading("九、阶段结论与后续建议", level=1)
    doc.add_paragraph(
        "2026-05-28 当日优化已形成可落地的推荐配置：在 Data05 验证集上，"
        "bg_y MAE 降至 0.0351 °/s（达标），bg_z MAE 为 0.0242 °/s（当前各轮最优），"
        "加速度计三轴与 bg_x 保持较好水平。相较 5 月 27 日基线，陀螺通道泛化问题得到明显缓解。"
    )
    doc.add_paragraph("后续建议：", style="List Bullet")
    for item in [
        "以 biasnet_weights_cmcc.weights.h5（165527 对应训练）作为当前主版本，接入 ekf_navigator（DeepBiasNet + window=300）。",
        "评估链路保持与训练一致：cmcc_stable 掩码 + 2 s EMA 后处理 + acc 校准 JSON。",
        "在新增场景数据上复训并交叉验证，确认 bg_y 稳态偏差是否进一步收敛。",
        "补充 EKF 端到端轨迹对比（有/无 BiasNet 补偿）并更新拒止段指标报表。",
    ]:
        doc.add_paragraph(item, style="List Bullet 2")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"已生成: {OUT_DOCX}")


if __name__ == "__main__":
    main()
